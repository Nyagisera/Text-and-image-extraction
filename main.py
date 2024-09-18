from flask import Flask, request, render_template, send_file
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import os
from PyPDF2 import PdfReader
import io
from PIL import Image
from textblob import TextBlob
from nltk.corpus import wordnet
import nltk
from symspellpy import SymSpell, Verbosity
import re

app = Flask(__name__)

# Configurations
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
IMAGE_FOLDER = 'images'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['IMAGE_FOLDER'] = IMAGE_FOLDER

# Initialize SymSpell
sym_spell = SymSpell(max_dictionary_edit_distance=2, prefix_length=7)
dictionary_path = "path_to_your_dictionary.txt"  # You need to provide this path
sym_spell.load_dictionary(dictionary_path, term_index=0, count_index=1)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files['file']
    search_term = request.form.get('search_term', '')
    search_type = request.form.get('search_type', '')

    if file:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        # Process the file and extract text/images
        extracted_text, images = extract_text_and_images(file_path)
        
        # Enhance search term
        search_term = correct_spelling(search_term)
        synonyms = get_synonyms(search_term)
        
        # Perform advanced search
        if search_type == 'word_length':
            filtered_text = search_by_word_length(extracted_text, int(search_term))
        elif search_type == 'numbers':
            filtered_text = extract_numbers(extracted_text)
        else:
            filtered_text = search_in_text(extracted_text, search_term, synonyms)
        
        return render_template('display_extracted.html', extracted_text=filtered_text, images=images, filename=file.filename)

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf_route():
    extracted_text = request.form.get('extracted_text')
    images_str = request.form.get('images')
    images = images_str.split(',') if images_str else []
    filename = request.form.get('filename')
    
    output_filename = f'extracted_{filename}.pdf'
    output_path = generate_pdf(extracted_text, images, output_filename)
    
    return send_file(output_path, as_attachment=True)

@app.route('/generate_docx', methods=['POST'])
def generate_docx_route():
    extracted_text = request.form.get('extracted_text')
    images_str = request.form.get('images')
    images = images_str.split(',') if images_str else []
    filename = request.form.get('filename')
    
    output_filename = f'extracted_{filename}.docx'
    output_path = generate_docx(extracted_text, images, output_filename)
    
    return send_file(output_path, as_attachment=True)

def extract_text_and_images(file_path):
    extracted_text = ""
    images = []

    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                extracted_text += page.extract_text() or ""

        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            data = xObject[obj]._data
                            image_filename = f"image_{i}_{obj[1:]}.png"
                            image_path = os.path.join(app.config['IMAGE_FOLDER'], image_filename)
                            with open(image_path, 'wb') as img_file:
                                img_file.write(data)
                            images.append(image_filename)
                except KeyError as e:
                    print(f"KeyError: {e} - This page might not contain images or has a different structure.")
                except Exception as e:
                    print(f"Exception: {e} - An error occurred while processing images.")

    except Exception as e:
        print(f"Exception: {e} - An error occurred while extracting text and images.")
    
    return extracted_text, images

def correct_spelling(term):
    # Correct spelling using TextBlob
    blob = TextBlob(term)
    corrected_term = str(blob.correct())
    
    # Correct spelling using SymSpell
    suggestions = sym_spell.lookup(corrected_term, Verbosity.CLOSEST, max_edit_distance=2)
    if suggestions:
        corrected_term = suggestions[0].term
    
    return corrected_term

def get_synonyms(word):
    synonyms = set()
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.add(lemma.name())
    return synonyms

def search_in_text(text, search_term, synonyms):
    # Convert synonyms set to a list
    synonym_list = list(synonyms)
    
    # Case-insensitive search for the term and its synonyms
    lines = text.split('\n')
    filtered_lines = []
    
    for line in lines:
        line_lower = line.lower()
        if search_term in line_lower or any(syn.lower() in line_lower for syn in synonym_list):
            filtered_lines.append(line)
    
    return '\n'.join(filtered_lines)

def search_by_word_length(text, length):
    words = re.findall(r'\b\w{' + str(length) + r'}\b', text)
    return '\n'.join(words)

def extract_numbers(text):
    numbers = re.findall(r'\b\d+\b', text)
    return '\n'.join(numbers)

def generate_pdf(content, images, output_filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font('Arial', '', 12)

    pdf.multi_cell(0, 10, content.encode('latin-1', 'replace').decode('latin-1'))

    if images:
        for image in images:
            pdf.add_page()
            image_path = os.path.join(app.config['IMAGE_FOLDER'], image)
            pdf.image(image_path, x=10, y=10, w=180)
    
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
    pdf.output(output_path)
    return output_path

def generate_docx(content, images, output_filename):
    doc = Document()
    doc.add_heading('Extracted Report', 0)
    doc.add_paragraph(content)

    if images:
        for image in images:
            image_path = os.path.join(app.config['IMAGE_FOLDER'], image)
            doc.add_page_break()
            doc.add_picture(image_path, width=Inches(6))
    
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    nltk.download('wordnet')
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(IMAGE_FOLDER, exist_ok=True)
    app.run(debug=True)
