from fpdf import FPDF

from docx import Document

def create_pdf(filename, content):
    """
    Creates a PDF report from the extracted content.
    
    Parameters:
    - filename: Name of the output PDF file.
    - content: Extracted text content.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    
    for line in content.split("\n"):
        pdf.multi_cell(0, 10, line)
    
    pdf.output(filename)


def create_docx(filename, content):
    """
    Creates a DOCX report from the extracted content.
    
    Parameters:
    - filename: Name of the output DOCX file.
    - content: Extracted text content.
    """
    doc = Document()
    doc.add_heading('Extracted Report', 0)

    for line in content.split("\n"):
        doc.add_paragraph(line)
    
    doc.save(filename)
